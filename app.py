import streamlit as st
import requests
import docx
from io import BytesIO
import re
import time

st.title("Corrector de Ortografía y Gramática")

# Leer la clave de API desde los secrets de Streamlit
API_KEY = st.secrets["openrouter_api_key"]

# Widget para subir el archivo
uploaded_file = st.file_uploader("Sube un documento de Word", type=["docx"])

if uploaded_file is not None:
    # Añadir un botón "Corregir"
    if st.button("Corregir"):
        # Leer el documento de Word
        document = docx.Document(uploaded_file)
        full_text = []
        for para in document.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)

        # Dividir párrafos muy largos
        max_paragraph_length = 500  # Puedes ajustar este valor según tus necesidades
        paragraphs = []

        def split_into_sentences(text):
            # Tokenizador simple de oraciones
            sentence_endings = re.compile(r'([.!?])\s+')
            sentences = sentence_endings.split(text)
            # Combinar fragmentos de oraciones
            combined = []
            for i in range(0, len(sentences)-1, 2):
                combined.append(sentences[i] + sentences[i+1])
            if len(sentences) % 2 == 1:
                combined.append(sentences[-1])
            return combined

        for para in text.split('\n'):
            if len(para) > max_paragraph_length:
                # Dividir el párrafo en oraciones
                sentences = split_into_sentences(para)
                temp_para = ''
                for sentence in sentences:
                    if len(temp_para) + len(sentence) <= max_paragraph_length:
                        temp_para += sentence + ' '
                    else:
                        paragraphs.append(temp_para.strip())
                        temp_para = sentence + ' '
                if temp_para:
                    paragraphs.append(temp_para.strip())
            else:
                paragraphs.append(para)

        # Unir los párrafos procesados
        processed_text = '\n'.join(paragraphs)

        # Mostrar barra de progreso
        progress_bar = st.progress(0)
        progress_text = st.empty()

        # Preparar la solicitud a la API enfocada en las nuevas especificaciones
        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }

        prompt = f"""
Corrige los errores de ortografía y gramática del siguiente texto, eliminando repeticiones innecesarias. Si durante la corrección se suprime una llamada a nota al pie, coloca un "+" en su lugar. No cambies el estilo, la redacción ni la estructura original. Mantén el significado y la forma del texto tal como están.

Texto a corregir:

{processed_text}
"""

        payload = {
            "model": "qwen/qwen-2.5-72b-instruct",
            "messages": [{"role": "user", "content": prompt}]
        }

        # Enviar la solicitud a la API de OpenRouter
        try:
            with st.spinner('Corrigiendo el documento...'):
                response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload, timeout=300)
                # Simular progreso mientras se procesa la solicitud
                for percent_complete in range(100):
                    time.sleep(0.05)  # Ajusta el tiempo según sea necesario
                    progress_bar.progress(percent_complete + 1)
        except requests.exceptions.Timeout:
            st.error("La solicitud ha excedido el tiempo de espera.")
        except Exception as e:
            st.error(f"Ocurrió un error: {e}")
        else:
            if response.status_code == 200:
                result = response.json()
                corrected_text = result["choices"][0]["message"]["content"]

                # Crear un nuevo documento de Word con el texto corregido
                corrected_document = docx.Document()
                for para in corrected_text.split('\n'):
                    corrected_document.add_paragraph(para)

                # Guardar en un objeto BytesIO
                corrected_io = BytesIO()
                corrected_document.save(corrected_io)
                corrected_io.seek(0)

                # Botón para descargar el documento corregido
                st.download_button(
                    label="Descargar Documento Corregido",
                    data=corrected_io,
                    file_name="documento_corregido.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error("Error en la solicitud a la API")
                st.error(response.text)
